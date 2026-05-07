"""
Carrier Analysis Tool
---------------------
Streamlit app for analyzing carrier tracking and milestone completeness data.

Workflow:
1. Upload an Excel/CSV export
2. Phase 1: Untracked shipment analysis with pattern detection
3. Phase 2: Milestone completeness analysis (tracked shipments only)
4. Log issues with evidence BOLs and screenshots throughout
5. Download a one-pager .docx for use with downstream email-drafting prompts
"""

import streamlit as st
import pandas as pd
import re
import io
import json
from datetime import datetime
from collections import Counter
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from PIL import Image

# ============================================================
# CONFIGURATION
# ============================================================

EXAMPLE_COUNT_UNTRACKED = 15
EXAMPLE_COUNT_MILESTONE = 12

# Pattern detection thresholds
LOCATION_PATTERN_MIN_SHARE = 0.30   # ≥30% of failures share a location
LOCATION_PATTERN_MIN_COUNT = 3       # at least 3 shipments
METHOD_PATTERN_MIN_SHARE = 0.50      # ≥50% of failures use one method
METHOD_PATTERN_MIN_COUNT = 3
ELD_PATTERN_MIN_COUNT = 3            # equipment ID seen ≥3 times in failures
LANE_PATTERN_MIN_SHARE = 0.30
LANE_PATTERN_MIN_COUNT = 3

# Reusable downstream prompt
EMAIL_PROMPT_TEMPLATE = """I've attached a one-pager that documents tracking and milestone completeness issues identified for a specific carrier. The document contains:
- A summary of overall tracking and milestone performance
- A list of specific issues found, each with a description, example BOLs, and screenshots where available

Please draft a professional but firm email to send to this carrier that:
1. Opens by stating the visibility/tracking concerns we've identified over the analysis period
2. Walks through each issue clearly, referencing the BOL examples and screenshots from the attached document
3. Asks for a specific action or response from the carrier on each issue
4. Closes by requesting a follow-up call or written response within a reasonable timeframe

Keep the tone constructive — we want to fix the relationship, not damage it — but be specific enough that the carrier cannot dismiss the concerns. Use the screenshots inline in the email where they support a claim."""

# Canonical column names → list of acceptable variants (lowercased, will be normalized)
COLUMN_ALIASES = {
    'tenant_name':              ['tenant name', 'tenant', 'tenantname'],
    'carrier_name':             ['carrier name', 'carrier', 'carriername'],
    'carrier_identifier':       ['carrier identifier selection', 'carrier identifier', 'identifier selection', 'identifier'],
    'scac':                     ['scac', 'scac value', 'scacs', 'scac values'],
    'bol':                      ['bill of lading', 'bol', 'billoflading', 'bill_of_lading', 'bol number'],
    'order_number':             ['order number', 'order_number', 'ordernumber', 'order'],
    'tracked':                  ['tracked', 'is tracked', 'is_tracked'],
    'tracking_type':            ['tracking type', 'trackingtype'],
    'tracking_method':          ['tracking method', 'trackingmethod'],
    'active_equipment_id':      ['active equipment id', 'active equipment', 'active_equipment_id', 'active eqp id'],
    'historical_equipment_id':  ['historical equipment id', 'historical equipment', 'historic equipment id', 'historic equipment'],
    'pickup_name':              ['pickup name', 'pickupname', 'origin name'],
    'pickup_city_state':        ['pickup city state', 'pickup city, state', 'pickup citystate', 'origin city state'],
    'pickup_country':           ['pickup country', 'origin country'],
    'pickup_region':            ['pickup region', 'origin region'],
    'dropoff_name':             ['drop-off name', 'dropoff name', 'drop off name', 'delivery name', 'destination name'],
    'dropoff_city_state':       ['drop-off city state', 'drop-off city, state', 'dropoff city state', 'drop off city, state', 'destination city state'],
    'dropoff_country':          ['drop-off country', 'dropoff country', 'destination country'],
    'dropoff_region':           ['drop-off region', 'dropoff region', 'destination region', 'dropoff country region', 'drop-off country region'],
    'final_status_reason':      ['final status reason', 'status reason'],
    'created_timestamp_date':   ['created timestamp date', 'created timestamp', 'created date'],
    'pickup_arrival':           ['pickup arrival utc timestamp raw', 'pickup arrival', 'pickup arrival utc', 'pickup arrival timestamp'],
    'pickup_departure':         ['pickup departure utc timestamp raw', 'pickup departure', 'pickup departure utc', 'pickup departure timestamp'],
    'dropoff_arrival':          ['drop-off arrival utc timestamp raw', 'dropoff arrival', 'drop off arrival utc', 'drop-off arrival', 'destination arrival'],
    'dropoff_departure':        ['drop-off departure utc timestamp raw', 'dropoff departure', 'drop off departure utc', 'drop-off departure', 'destination departure'],
}

# ============================================================
# UTILITY FUNCTIONS
# ============================================================

def normalize(s):
    """Lowercase, strip non-alphanumeric, collapse whitespace, and canonicalize
    common compound words ('drop off'/'drop-off'/'dropoff' all → 'dropoff';
    'pick up'/'pick-up'/'pickup' all → 'pickup')."""
    if pd.isna(s):
        return ''
    s = str(s).lower().strip()
    s = re.sub(r'[^a-z0-9]+', ' ', s)
    s = re.sub(r'\s+', ' ', s).strip()
    # Canonicalize compound words so spelling variants collapse together
    s = re.sub(r'\bdrop\s?off\b', 'dropoff', s)
    s = re.sub(r'\bpick\s?up\b', 'pickup', s)
    return s


def map_columns(df):
    """Map actual columns in the dataframe to canonical names via fuzzy matching."""
    actual_cols = {normalize(c): c for c in df.columns}
    mapping = {}
    for canonical, aliases in COLUMN_ALIASES.items():
        for alias in aliases:
            n = normalize(alias)
            if n in actual_cols:
                mapping[canonical] = actual_cols[n]
                break
    return mapping


def parse_bool(val):
    """Parse various truthy/falsy representations."""
    if pd.isna(val):
        return None
    s = str(val).strip().lower()
    if s in ('true', 't', 'yes', 'y', '1'):
        return True
    if s in ('false', 'f', 'no', 'n', '0'):
        return False
    return None


def is_missing(val):
    """Check whether a timestamp value is missing/blank."""
    if pd.isna(val):
        return True
    s = str(val).strip()
    if s == '' or s.lower() in ('nan', 'none', 'null'):
        return True
    return False


def safe_get(row, col_map, key, default=''):
    """Safely get a value from a row using the column mapping."""
    actual_col = col_map.get(key)
    if actual_col is None or actual_col not in row.index:
        return default
    val = row[actual_col]
    if pd.isna(val):
        return default
    return val


def load_file(uploaded_file):
    """Load Excel or CSV. Returns a DataFrame."""
    name = uploaded_file.name.lower()
    if name.endswith('.csv'):
        return pd.read_csv(uploaded_file, dtype=str, keep_default_na=True)
    elif name.endswith('.xlsx') or name.endswith('.xlsm') or name.endswith('.xls'):
        return pd.read_excel(uploaded_file, dtype=str, engine='openpyxl' if name.endswith('xlsx') or name.endswith('xlsm') else None)
    else:
        raise ValueError("Unsupported file type. Please upload .csv or .xlsx")


# ============================================================
# PATTERN DETECTION
# ============================================================

def detect_location_concentration(failures_df, all_df, col, label, min_share, min_count):
    """Detect if failures concentrate at a single location, AND that location's
    failure rate exceeds the overall failure rate."""
    if col not in failures_df.columns or len(failures_df) < min_count:
        return None

    counts = failures_df[col].fillna('(blank)').replace('', '(blank)').value_counts()
    if counts.empty:
        return None

    top_value = counts.index[0]
    top_count = int(counts.iloc[0])

    if top_count < min_count:
        return None

    share = top_count / len(failures_df)
    if share < min_share:
        return None

    # Compare failure rate at this location vs overall
    location_total = int(((all_df[col].fillna('(blank)').replace('', '(blank)')) == top_value).sum())
    overall_failure_rate = len(failures_df) / len(all_df) if len(all_df) > 0 else 0
    location_failure_rate = top_count / location_total if location_total > 0 else 0

    # Only flag if the location is also a higher-failure-rate location than baseline
    if location_failure_rate <= overall_failure_rate * 1.2:
        return None

    return {
        'type': 'location',
        'dimension': label,
        'value': str(top_value),
        'count': top_count,
        'share_of_failures': share,
        'location_failure_rate': location_failure_rate,
        'overall_failure_rate': overall_failure_rate,
        'location_total': location_total,
        'message': (
            f"**{label}:** {top_count} of {len(failures_df)} failures "
            f"({share*100:.0f}%) are at '{top_value}'. "
            f"Failure rate at this {label.lower()} is {location_failure_rate*100:.0f}% "
            f"vs overall {overall_failure_rate*100:.0f}%."
        )
    }


def detect_lane_concentration(failures_df, all_df, pickup_col, dropoff_col, min_share, min_count):
    """Detect if failures concentrate on a specific origin→destination lane."""
    if pickup_col not in failures_df.columns or dropoff_col not in failures_df.columns:
        return None
    if len(failures_df) < min_count:
        return None

    failures_df = failures_df.copy()
    failures_df['_lane'] = (
        failures_df[pickup_col].fillna('(blank)').astype(str) + ' → ' +
        failures_df[dropoff_col].fillna('(blank)').astype(str)
    )
    all_df = all_df.copy()
    all_df['_lane'] = (
        all_df[pickup_col].fillna('(blank)').astype(str) + ' → ' +
        all_df[dropoff_col].fillna('(blank)').astype(str)
    )

    counts = failures_df['_lane'].value_counts()
    if counts.empty:
        return None

    top_lane = counts.index[0]
    top_count = int(counts.iloc[0])

    if top_count < min_count:
        return None
    share = top_count / len(failures_df)
    if share < min_share:
        return None

    lane_total = int((all_df['_lane'] == top_lane).sum())
    overall_failure_rate = len(failures_df) / len(all_df) if len(all_df) > 0 else 0
    lane_failure_rate = top_count / lane_total if lane_total > 0 else 0

    if lane_failure_rate <= overall_failure_rate * 1.2:
        return None

    return {
        'type': 'lane',
        'value': top_lane,
        'count': top_count,
        'share_of_failures': share,
        'lane_failure_rate': lane_failure_rate,
        'overall_failure_rate': overall_failure_rate,
        'message': (
            f"**Lane:** {top_count} of {len(failures_df)} failures "
            f"({share*100:.0f}%) are on the lane '{top_lane}'. "
            f"Failure rate on this lane is {lane_failure_rate*100:.0f}% "
            f"vs overall {overall_failure_rate*100:.0f}%."
        )
    }


def detect_tracking_method_concentration(failures_df, all_df, col_map, min_share, min_count):
    """If multiple tracking methods exist in the file, flag if failures
    concentrate disproportionately on one method."""
    method_col = col_map.get('tracking_method')
    if not method_col or method_col not in all_df.columns:
        return None

    all_methods = all_df[method_col].dropna().astype(str)
    all_methods = all_methods[all_methods.str.strip() != '']
    unique_methods = all_methods.unique()

    if len(unique_methods) < 2:
        return None  # Only one method in use, no signal here
    if len(failures_df) < min_count:
        return None

    failure_methods = failures_df[method_col].dropna().astype(str)
    failure_methods = failure_methods[failure_methods.str.strip() != '']

    if failure_methods.empty:
        return None

    counts = failure_methods.value_counts()
    top_method = counts.index[0]
    top_count = int(counts.iloc[0])

    if top_count < min_count:
        return None
    share_of_failures = top_count / len(failure_methods)
    if share_of_failures < min_share:
        return None

    # Compare to method's share of overall volume
    method_total = int((all_methods == top_method).sum())
    method_share_of_total = method_total / len(all_methods) if len(all_methods) > 0 else 0

    # Only flag if failures are over-represented on this method vs its overall share
    if share_of_failures <= method_share_of_total + 0.1:
        return None

    return {
        'type': 'tracking_method',
        'value': top_method,
        'count': top_count,
        'share_of_failures': share_of_failures,
        'method_share_of_total': method_share_of_total,
        'message': (
            f"**Tracking method:** {top_count} of {len(failure_methods)} failures "
            f"({share_of_failures*100:.0f}%) use tracking method '{top_method}', "
            f"but this method only accounts for {method_share_of_total*100:.0f}% "
            f"of all shipments. Disproportionate failure concentration."
        )
    }


def detect_eld_equipment_concentration(failures_df, col_map, min_count):
    """If tracking method is ELD, flag if a specific active equipment ID
    appears repeatedly in failures."""
    method_col = col_map.get('tracking_method')
    eqp_col = col_map.get('active_equipment_id')
    if not method_col or not eqp_col:
        return None
    if method_col not in failures_df.columns or eqp_col not in failures_df.columns:
        return None

    eld_failures = failures_df[
        failures_df[method_col].astype(str).str.lower().str.contains('eld', na=False)
    ]
    if len(eld_failures) < min_count:
        return None

    eqp = eld_failures[eqp_col].dropna().astype(str)
    eqp = eqp[eqp.str.strip() != '']
    if eqp.empty:
        return None

    counts = eqp.value_counts()
    top_eqp = counts.index[0]
    top_count = int(counts.iloc[0])

    if top_count < min_count:
        return None

    return {
        'type': 'eld_equipment',
        'value': top_eqp,
        'count': top_count,
        'total_eld_failures': len(eld_failures),
        'message': (
            f"**ELD equipment ID:** Active equipment ID '{top_eqp}' appears in "
            f"{top_count} of {len(eld_failures)} ELD-tracked failures. "
            f"Possible faulty/misconfigured device."
        )
    }


# ============================================================
# PHASE 1: UNTRACKED ANALYSIS
# ============================================================

def analyze_untracked(df, col_map):
    """Compute untracked stats, examples, and patterns."""
    tracked_col = col_map.get('tracked')
    if not tracked_col:
        return None

    df = df.copy()
    df['_tracked_bool'] = df[tracked_col].apply(parse_bool)

    # Drop rows where we can't parse tracked status
    valid = df[df['_tracked_bool'].notna()]
    untracked = valid[valid['_tracked_bool'] == False]
    tracked = valid[valid['_tracked_bool'] == True]

    total = len(valid)
    n_unt = len(untracked)
    n_tr = len(tracked)

    result = {
        'total': total,
        'tracked_count': n_tr,
        'untracked_count': n_unt,
        'tracked_pct': (n_tr / total * 100) if total else 0,
        'untracked_pct': (n_unt / total * 100) if total else 0,
        'untracked_df': untracked,
        'tracked_df': tracked,
        'patterns': [],
    }

    # Pattern detection on untracked vs full dataset
    patterns = []
    for key, label in [('pickup_name', 'Pickup location'),
                       ('dropoff_name', 'Drop-off location'),
                       ('pickup_city_state', 'Pickup city/state'),
                       ('dropoff_city_state', 'Drop-off city/state')]:
        col = col_map.get(key)
        if col:
            p = detect_location_concentration(
                untracked, valid, col, label,
                LOCATION_PATTERN_MIN_SHARE, LOCATION_PATTERN_MIN_COUNT
            )
            if p:
                patterns.append(p)

    # Lane pattern
    pn = col_map.get('pickup_name')
    dn = col_map.get('dropoff_name')
    if pn and dn:
        p = detect_lane_concentration(untracked, valid, pn, dn,
                                      LANE_PATTERN_MIN_SHARE, LANE_PATTERN_MIN_COUNT)
        if p:
            patterns.append(p)

    # Tracking method
    p = detect_tracking_method_concentration(
        untracked, valid, col_map,
        METHOD_PATTERN_MIN_SHARE, METHOD_PATTERN_MIN_COUNT
    )
    if p:
        patterns.append(p)

    # ELD equipment ID
    p = detect_eld_equipment_concentration(untracked, col_map, ELD_PATTERN_MIN_COUNT)
    if p:
        patterns.append(p)

    result['patterns'] = patterns
    return result


# ============================================================
# PHASE 2: MILESTONE COMPLETENESS (TRACKED ONLY)
# ============================================================

def compute_milestone_buckets(tracked_df, col_map):
    """Annotate tracked_df with missing-milestone flags and group into the 16 buckets."""
    pa = col_map.get('pickup_arrival')
    pd_ = col_map.get('pickup_departure')
    da = col_map.get('dropoff_arrival')
    dd = col_map.get('dropoff_departure')

    df = tracked_df.copy()
    df['_pa_missing'] = df[pa].apply(is_missing) if pa else True
    df['_pd_missing'] = df[pd_].apply(is_missing) if pd_ else True
    df['_da_missing'] = df[da].apply(is_missing) if da else True
    df['_dd_missing'] = df[dd].apply(is_missing) if dd else True

    df['_pattern'] = list(zip(df['_pa_missing'], df['_pd_missing'],
                               df['_da_missing'], df['_dd_missing']))

    # Individual missing rates
    individual = {
        'Pickup Arrival':    df['_pa_missing'].sum() / len(df) if len(df) else 0,
        'Pickup Departure':  df['_pd_missing'].sum() / len(df) if len(df) else 0,
        'Dropoff Arrival':   df['_da_missing'].sum() / len(df) if len(df) else 0,
        'Dropoff Departure': df['_dd_missing'].sum() / len(df) if len(df) else 0,
    }

    # Bucket counts
    bucket_counts = df['_pattern'].value_counts().to_dict()
    buckets = []
    for pattern, count in bucket_counts.items():
        subset = df[df['_pattern'] == pattern]
        buckets.append({
            'pattern': pattern,
            'label': pattern_label(pattern),
            'count': int(count),
            'data': subset,
        })

    # Sort by count descending, with "complete" pinned to bottom
    def sort_key(b):
        is_complete = (b['pattern'] == (False, False, False, False))
        # Complete goes last regardless of count
        return (1 if is_complete else 0, -b['count'])
    buckets.sort(key=sort_key)

    return df, buckets, individual


def pattern_label(pattern):
    pa, pd_, da, dd = pattern
    missing = []
    if pa: missing.append("Pickup Arrival")
    if pd_: missing.append("Pickup Departure")
    if da: missing.append("Dropoff Arrival")
    if dd: missing.append("Dropoff Departure")
    if not missing:
        return "✅ All 4 milestones present (Complete)"
    if len(missing) == 4:
        return "❌ All 4 milestones missing"
    return "Missing: " + " + ".join(missing)


def analyze_milestone_bucket(bucket_df, all_tracked_df, col_map):
    """Run pattern detection on a specific bucket."""
    patterns = []
    for key, label in [('pickup_name', 'Pickup location'),
                       ('dropoff_name', 'Drop-off location')]:
        col = col_map.get(key)
        if col:
            p = detect_location_concentration(
                bucket_df, all_tracked_df, col, label,
                LOCATION_PATTERN_MIN_SHARE, LOCATION_PATTERN_MIN_COUNT
            )
            if p:
                patterns.append(p)

    pn = col_map.get('pickup_name')
    dn = col_map.get('dropoff_name')
    if pn and dn:
        p = detect_lane_concentration(bucket_df, all_tracked_df, pn, dn,
                                      LANE_PATTERN_MIN_SHARE, LANE_PATTERN_MIN_COUNT)
        if p:
            patterns.append(p)

    p = detect_tracking_method_concentration(
        bucket_df, all_tracked_df, col_map,
        METHOD_PATTERN_MIN_SHARE, METHOD_PATTERN_MIN_COUNT
    )
    if p:
        patterns.append(p)

    p = detect_eld_equipment_concentration(bucket_df, col_map, ELD_PATTERN_MIN_COUNT)
    if p:
        patterns.append(p)

    return patterns


# ============================================================
# DOCX EXPORT
# ============================================================

def build_docx(state, col_map):
    """Build the one-pager .docx report from session state."""
    doc = Document()

    # Header
    carrier_name = state.get('carrier_name', 'Unknown Carrier')
    tenant_name = state.get('tenant_name', 'Unknown Tenant')

    title = doc.add_heading(f'Carrier Analysis: {carrier_name}', 0)

    subtitle = doc.add_paragraph()
    subtitle.add_run(f'Tenant: {tenant_name}\n').bold = True
    subtitle.add_run(f'Generated: {datetime.now().strftime("%Y-%m-%d %H:%M")}\n')

    # Summary stats
    doc.add_heading('Performance Summary', 1)
    summary = state.get('untracked_summary', {})
    if summary:
        p = doc.add_paragraph()
        p.add_run(f"Total shipments analyzed: {summary.get('total', 0)}\n").bold = True
        p.add_run(f"Tracked: {summary.get('tracked_count', 0)} ({summary.get('tracked_pct', 0):.1f}%)\n")
        p.add_run(f"Untracked: {summary.get('untracked_count', 0)} ({summary.get('untracked_pct', 0):.1f}%)\n")

    ms = state.get('milestone_summary', {})
    if ms:
        p = doc.add_paragraph()
        p.add_run(f"\nMilestone completeness rate: {ms.get('completeness_pct', 0):.1f}% "
                  f"of {ms.get('tracked_total', 0)} tracked shipments\n").bold = True
        for name, rate in ms.get('individual', {}).items():
            p.add_run(f"  • {name} present: {(1-rate)*100:.1f}%\n")

    # Tracking issues
    tracking_issues = state.get('tracking_issues', [])
    if tracking_issues:
        doc.add_heading('Tracking Issues (Untracked Shipments)', 1)
        for i, issue in enumerate(tracking_issues, 1):
            doc.add_heading(f"Issue {i}: {issue.get('title', 'Untitled')}", 2)
            doc.add_paragraph(issue.get('description', ''))

            bols = issue.get('bols', [])
            if bols:
                p = doc.add_paragraph()
                p.add_run('Example BOLs: ').bold = True
                p.add_run(', '.join(bols))

            for img_bytes, img_name in issue.get('screenshots', []):
                try:
                    img_stream = io.BytesIO(img_bytes)
                    doc.add_picture(img_stream, width=Inches(6))
                    cap = doc.add_paragraph(img_name)
                    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    cap.runs[0].italic = True
                    cap.runs[0].font.size = Pt(9)
                except Exception as e:
                    doc.add_paragraph(f"[Could not embed image: {img_name}]")

    # Milestone issues
    milestone_issues = state.get('milestone_issues', [])
    if milestone_issues:
        doc.add_heading('Milestone Completeness Issues', 1)
        for i, issue in enumerate(milestone_issues, 1):
            doc.add_heading(f"Issue {i}: {issue.get('title', 'Untitled')}", 2)
            if issue.get('bucket_label'):
                p = doc.add_paragraph()
                p.add_run('Bucket: ').bold = True
                p.add_run(issue['bucket_label'])
            doc.add_paragraph(issue.get('description', ''))

            bols = issue.get('bols', [])
            if bols:
                p = doc.add_paragraph()
                p.add_run('Example BOLs: ').bold = True
                p.add_run(', '.join(bols))

            for img_bytes, img_name in issue.get('screenshots', []):
                try:
                    img_stream = io.BytesIO(img_bytes)
                    doc.add_picture(img_stream, width=Inches(6))
                    cap = doc.add_paragraph(img_name)
                    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    cap.runs[0].italic = True
                    cap.runs[0].font.size = Pt(9)
                except Exception:
                    doc.add_paragraph(f"[Could not embed image: {img_name}]")

    if not tracking_issues and not milestone_issues:
        doc.add_paragraph('No issues were logged during this analysis.')

    # Save to bytes
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


def build_progress_json(state):
    """Serialize session state (without screenshots) for save/restore."""
    out = {
        'carrier_name': state.get('carrier_name'),
        'tenant_name': state.get('tenant_name'),
        'tracking_issues': [
            {
                'title': i.get('title'),
                'description': i.get('description'),
                'bols': i.get('bols', []),
                # Screenshots are NOT serialized (binary)
            }
            for i in state.get('tracking_issues', [])
        ],
        'milestone_issues': [
            {
                'title': i.get('title'),
                'description': i.get('description'),
                'bols': i.get('bols', []),
                'bucket_label': i.get('bucket_label'),
            }
            for i in state.get('milestone_issues', [])
        ],
        'phase': state.get('phase', 'tracking'),
    }
    return json.dumps(out, indent=2).encode('utf-8')


# ============================================================
# UI HELPERS
# ============================================================

def render_examples_table(df, col_map, key_prefix=''):
    """Render a small table of example shipments showing BOL + key context."""
    cols_to_show = []
    headers = []
    for key, header in [
        ('bol', 'BOL'),
        ('order_number', 'Order #'),
        ('pickup_name', 'Pickup'),
        ('dropoff_name', 'Drop-off'),
        ('tracking_method', 'Method'),
        ('active_equipment_id', 'Active Eqp ID'),
        ('final_status_reason', 'Status Reason'),
    ]:
        col = col_map.get(key)
        if col and col in df.columns:
            cols_to_show.append(col)
            headers.append(header)

    if not cols_to_show:
        st.info("No displayable columns found.")
        return

    display_df = df[cols_to_show].copy()
    display_df.columns = headers
    display_df = display_df.fillna('—')
    st.dataframe(display_df, use_container_width=True, hide_index=True)


def issue_logger(phase_key, bucket_label=None):
    """Render the issue logging widget. Issues stored in state[phase_key+'_issues']."""
    state_key = f'{phase_key}_issues'
    if state_key not in st.session_state:
        st.session_state[state_key] = []

    # Display existing issues
    issues = st.session_state[state_key]
    if issues:
        st.markdown("##### Logged issues")
        for i, issue in enumerate(issues):
            with st.expander(f"Issue {i+1}: {issue.get('title') or '(untitled)'}", expanded=False):
                st.write(f"**Description:** {issue.get('description', '')}")
                if issue.get('bucket_label'):
                    st.write(f"**Bucket:** {issue['bucket_label']}")
                if issue.get('bols'):
                    st.write(f"**BOLs:** {', '.join(issue['bols'])}")
                if issue.get('screenshots'):
                    for img_bytes, img_name in issue['screenshots']:
                        st.image(img_bytes, caption=img_name, width=400)
                if st.button("Delete this issue", key=f"del_{phase_key}_{i}"):
                    st.session_state[state_key].pop(i)
                    st.rerun()

    # New issue form
    with st.expander("➕ Log a new issue", expanded=False):
        with st.form(key=f'{phase_key}_form_{len(issues)}', clear_on_submit=True):
            title = st.text_input("Issue title (short)", key=f'title_{phase_key}')
            description = st.text_area("Description (what's the problem, what's likely the cause)",
                                        height=120, key=f'desc_{phase_key}')
            bols_raw = st.text_area("Evidence BOLs (one per line, or comma-separated)",
                                     height=80, key=f'bols_{phase_key}')
            screenshots = st.file_uploader(
                "Screenshots (optional, multiple allowed)",
                type=['png', 'jpg', 'jpeg'],
                accept_multiple_files=True,
                key=f'shots_{phase_key}'
            )
            submitted = st.form_submit_button("Save issue")

            if submitted:
                if not title and not description:
                    st.error("Add at least a title or description.")
                else:
                    bols = [b.strip() for b in re.split(r'[,\n]', bols_raw) if b.strip()]
                    shot_data = []
                    for s in (screenshots or []):
                        shot_data.append((s.read(), s.name))
                    new_issue = {
                        'title': title.strip(),
                        'description': description.strip(),
                        'bols': bols,
                        'screenshots': shot_data,
                    }
                    if bucket_label:
                        new_issue['bucket_label'] = bucket_label
                    st.session_state[state_key].append(new_issue)
                    st.success("Issue saved.")
                    st.rerun()


# ============================================================
# MAIN APP
# ============================================================

def main():
    st.set_page_config(page_title="Carrier Analysis Tool", layout="wide")

    # Initialize session state
    defaults = {
        'phase': 'upload',
        'df': None,
        'col_map': None,
        'untracked_summary': None,
        'milestone_summary': None,
        'tracking_issues': [],
        'milestone_issues': [],
        'carrier_name': None,
        'tenant_name': None,
    }
    for k, v in defaults.items():
        if k not in st.session_state:
            st.session_state[k] = v

    st.title("📋 Carrier Analysis Tool")
    st.caption("Upload a shipment export and walk through tracking + milestone completeness analysis.")

    # ---------- Sidebar: navigation + progress save/load ----------
    with st.sidebar:
        st.markdown("### Navigation")
        phase_labels = {
            'upload':    '1. Upload',
            'tracking':  '2. Tracking analysis',
            'milestone': '3. Milestone completeness',
            'export':    '4. Export & email prompt',
        }
        for p, label in phase_labels.items():
            indicator = "▶ " if st.session_state.phase == p else "  "
            st.write(f"{indicator}{label}")

        st.markdown("---")
        st.markdown("### Progress")
        st.write(f"Tracking issues logged: **{len(st.session_state.tracking_issues)}**")
        st.write(f"Milestone issues logged: **{len(st.session_state.milestone_issues)}**")

        if st.session_state.df is not None:
            st.markdown("---")
            st.download_button(
                "💾 Save progress (JSON)",
                data=build_progress_json(st.session_state),
                file_name=f"carrier_analysis_progress_{datetime.now().strftime('%Y%m%d_%H%M')}.json",
                mime="application/json",
                help="Download a snapshot of your logged issues. Screenshots are NOT included — only text."
            )

            uploaded_progress = st.file_uploader("Restore progress from JSON", type=['json'],
                                                  key='progress_restore')
            if uploaded_progress:
                try:
                    data = json.loads(uploaded_progress.read())
                    st.session_state.tracking_issues = [
                        {**i, 'screenshots': []} for i in data.get('tracking_issues', [])
                    ]
                    st.session_state.milestone_issues = [
                        {**i, 'screenshots': []} for i in data.get('milestone_issues', [])
                    ]
                    st.success("Progress restored. (Screenshots not restored — re-upload if needed.)")
                except Exception as e:
                    st.error(f"Could not parse progress file: {e}")

        st.markdown("---")
        if st.button("🔄 Reset everything"):
            for k in list(st.session_state.keys()):
                del st.session_state[k]
            st.rerun()

    # ---------- PHASE: UPLOAD ----------
    if st.session_state.phase == 'upload':
        st.header("Upload shipment export")
        uploaded = st.file_uploader("Excel (.xlsx) or CSV", type=['xlsx', 'csv', 'xlsm'])

        if uploaded:
            try:
                df = load_file(uploaded)
            except Exception as e:
                st.error(f"Could not read the file: {e}")
                return

            col_map = map_columns(df)

            # Required columns
            required = ['tracked', 'bol']
            missing = [r for r in required if r not in col_map]
            if missing:
                st.error(f"Could not find required columns: {missing}. "
                         f"Found columns: {list(df.columns)}")
                return

            # Multi-carrier handling
            carrier_col = col_map.get('carrier_name')
            tenant_col = col_map.get('tenant_name')

            if carrier_col:
                carriers = df[carrier_col].dropna().unique()
                carriers = [c for c in carriers if str(c).strip() != '']
                if len(carriers) > 1:
                    st.warning(f"Multiple carriers detected in this file: {list(carriers)}. "
                               "Pick one to analyze.")
                    chosen = st.selectbox("Carrier to analyze", carriers, key='carrier_pick')
                    df = df[df[carrier_col] == chosen]
                    st.session_state.carrier_name = chosen
                elif len(carriers) == 1:
                    st.session_state.carrier_name = str(carriers[0])
                else:
                    st.session_state.carrier_name = "Unknown Carrier"

            if tenant_col:
                tenants = df[tenant_col].dropna().unique()
                if len(tenants) >= 1:
                    st.session_state.tenant_name = str(tenants[0])

            # Show detected columns for transparency
            with st.expander("Detected column mapping (click to verify)", expanded=False):
                map_df = pd.DataFrame([
                    {'Field': k, 'Mapped to column': v}
                    for k, v in col_map.items()
                ])
                st.dataframe(map_df, use_container_width=True, hide_index=True)

                missing_optional = [k for k in COLUMN_ALIASES if k not in col_map]
                if missing_optional:
                    st.caption(f"⚠️ Canonical fields not found in file: {', '.join(missing_optional)}")

                mapped_actual = set(col_map.values())
                unmapped_file_cols = [c for c in df.columns if c not in mapped_actual]
                if unmapped_file_cols:
                    st.caption(f"ℹ️ File columns not used by analysis (e.g. attributes, latency stats): {', '.join(unmapped_file_cols)}")

            st.success(f"Loaded {len(df)} rows. "
                       f"Carrier: **{st.session_state.carrier_name}**, "
                       f"Tenant: **{st.session_state.tenant_name}**")

            if st.button("Continue to tracking analysis →", type='primary'):
                st.session_state.df = df
                st.session_state.col_map = col_map
                st.session_state.phase = 'tracking'
                st.rerun()

    # ---------- PHASE: TRACKING ----------
    elif st.session_state.phase == 'tracking':
        df = st.session_state.df
        col_map = st.session_state.col_map

        st.header("Phase 1: Tracking Analysis")

        summary = analyze_untracked(df, col_map)
        if summary is None:
            st.error("Could not run analysis — 'tracked' column not found.")
            return
        st.session_state.untracked_summary = {
            k: v for k, v in summary.items() if k not in ('untracked_df', 'tracked_df', 'patterns')
        }

        # Headline split
        col1, col2, col3 = st.columns(3)
        col1.metric("Total shipments", summary['total'])
        col2.metric("Tracked", f"{summary['tracked_count']} ({summary['tracked_pct']:.1f}%)")
        col3.metric("Untracked", f"{summary['untracked_count']} ({summary['untracked_pct']:.1f}%)")

        st.markdown("---")

        # Patterns
        st.subheader("🔍 Detected patterns in untracked shipments")
        if summary['patterns']:
            for p in summary['patterns']:
                st.info(p['message'])
        else:
            st.caption("No strong patterns detected. Untracked shipments appear distributed across "
                       "locations, lanes, and tracking methods.")

        st.markdown("---")

        # Examples
        st.subheader(f"📋 Untracked examples (up to {EXAMPLE_COUNT_UNTRACKED})")
        examples = summary['untracked_df'].head(EXAMPLE_COUNT_UNTRACKED)
        if len(examples) == 0:
            st.success("No untracked shipments. 🎉")
        else:
            render_examples_table(examples, col_map, 'untracked')
            st.caption("Copy a BOL → paste into Movement to investigate.")

        st.markdown("---")

        # Issue logger
        st.subheader("📝 Log tracking issues")
        issue_logger('tracking')

        st.markdown("---")

        # Next phase button
        col_a, col_b = st.columns([1, 1])
        with col_a:
            if st.button("← Back to upload"):
                st.session_state.phase = 'upload'
                st.rerun()
        with col_b:
            if st.button("Next: Milestone completeness →", type='primary'):
                st.session_state.phase = 'milestone'
                st.rerun()

    # ---------- PHASE: MILESTONE ----------
    elif st.session_state.phase == 'milestone':
        df = st.session_state.df
        col_map = st.session_state.col_map

        st.header("Phase 2: Milestone Completeness")
        st.caption("Analysis below covers **tracked shipments only**.")

        # Get tracked subset
        tracked_col = col_map.get('tracked')
        df_clean = df.copy()
        df_clean['_tracked_bool'] = df_clean[tracked_col].apply(parse_bool)
        tracked_df = df_clean[df_clean['_tracked_bool'] == True]

        if len(tracked_df) == 0:
            st.warning("No tracked shipments in this dataset. Skip to export.")
            if st.button("Skip to export →", type='primary'):
                st.session_state.phase = 'export'
                st.rerun()
            return

        # Required milestone columns?
        ms_cols = ['pickup_arrival', 'pickup_departure', 'dropoff_arrival', 'dropoff_departure']
        missing_ms = [c for c in ms_cols if c not in col_map]
        if missing_ms:
            st.error(f"Missing required milestone timestamp columns: {missing_ms}. "
                     "Cannot perform milestone completeness analysis.")
            return

        annotated_df, buckets, individual = compute_milestone_buckets(tracked_df, col_map)

        # Headline metrics
        complete_count = sum(b['count'] for b in buckets if b['pattern'] == (False, False, False, False))
        completeness_pct = (complete_count / len(tracked_df) * 100) if len(tracked_df) else 0

        st.session_state.milestone_summary = {
            'tracked_total': len(tracked_df),
            'completeness_pct': completeness_pct,
            'individual': individual,
        }

        c1, c2 = st.columns(2)
        c1.metric("Tracked shipments analyzed", len(tracked_df))
        c2.metric("Milestone completeness", f"{completeness_pct:.1f}%",
                  help="Shipments with all 4 milestones present, as % of tracked shipments.")

        # Individual milestone presence rates
        st.markdown("##### Individual milestone presence rates")
        rate_df = pd.DataFrame([
            {'Milestone': name, 'Present %': f"{(1-rate)*100:.1f}%",
             'Missing count': int(rate * len(tracked_df))}
            for name, rate in individual.items()
        ])
        st.dataframe(rate_df, use_container_width=True, hide_index=True)

        st.markdown("---")

        # Bucket breakdown
        st.subheader("📊 Combination buckets (15 missing-patterns + complete), ranked by count")
        st.caption("Click a bucket to see examples and pattern hints.")

        for i, b in enumerate(buckets):
            if b['count'] == 0:
                continue
            is_complete = (b['pattern'] == (False, False, False, False))
            label = b['label']
            with st.expander(f"{label} — **{b['count']}** shipments", expanded=(i == 0 and not is_complete)):
                if is_complete:
                    st.success("These shipments have all 4 milestones. No issue to investigate here.")
                    continue

                # Pattern hints for this bucket
                bucket_patterns = analyze_milestone_bucket(b['data'], tracked_df, col_map)
                if bucket_patterns:
                    st.markdown("**Pattern hints for this bucket:**")
                    for p in bucket_patterns:
                        st.info(p['message'])
                else:
                    st.caption("No strong sub-pattern in this bucket.")

                # Examples
                st.markdown(f"**Examples (up to {EXAMPLE_COUNT_MILESTONE}):**")
                ex = b['data'].head(EXAMPLE_COUNT_MILESTONE)
                render_examples_table(ex, col_map, f'bucket_{i}')

        st.markdown("---")

        # Issue logger
        st.subheader("📝 Log milestone completeness issues")
        st.caption("Tip: when logging an issue, paste in 2–3 BOLs from the relevant bucket as evidence.")
        issue_logger('milestone')

        st.markdown("---")

        # Navigation
        col_a, col_b = st.columns([1, 1])
        with col_a:
            if st.button("← Back to tracking"):
                st.session_state.phase = 'tracking'
                st.rerun()
        with col_b:
            if st.button("Generate final notes →", type='primary'):
                st.session_state.phase = 'export'
                st.rerun()

    # ---------- PHASE: EXPORT ----------
    elif st.session_state.phase == 'export':
        st.header("Final notes & email prompt")

        st.subheader("📄 Download one-pager")
        st.write("Carrier-specific document with all logged issues, BOLs, and screenshots embedded.")

        carrier = st.session_state.get('carrier_name') or 'carrier'
        safe_name = re.sub(r'[^a-zA-Z0-9_-]+', '_', str(carrier))
        filename = f"carrier_analysis_{safe_name}_{datetime.now().strftime('%Y%m%d')}.docx"

        try:
            buf = build_docx(st.session_state, st.session_state.col_map or {})
            st.download_button(
                "📥 Download one-pager (.docx)",
                data=buf,
                file_name=filename,
                mime='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                type='primary'
            )
        except Exception as e:
            st.error(f"Could not build document: {e}")

        st.markdown("---")

        st.subheader("📋 Reusable email-drafting prompt")
        st.write("Copy this prompt and feed it (along with the .docx above) to a new Claude chat to draft the carrier email.")
        st.code(EMAIL_PROMPT_TEMPLATE, language='text')
        st.caption("This prompt stays the same for every carrier — only the attached document changes.")

        st.markdown("---")

        # Quick summary view
        st.subheader("Summary of logged issues")
        ti = st.session_state.tracking_issues
        mi = st.session_state.milestone_issues
        if not ti and not mi:
            st.info("No issues were logged during this analysis.")
        else:
            if ti:
                st.markdown(f"**Tracking issues ({len(ti)}):**")
                for i, x in enumerate(ti, 1):
                    st.write(f"{i}. {x.get('title') or '(untitled)'}")
            if mi:
                st.markdown(f"**Milestone issues ({len(mi)}):**")
                for i, x in enumerate(mi, 1):
                    label = f" [{x['bucket_label']}]" if x.get('bucket_label') else ''
                    st.write(f"{i}. {x.get('title') or '(untitled)'}{label}")

        st.markdown("---")
        col_a, col_b = st.columns(2)
        with col_a:
            if st.button("← Back to milestone phase"):
                st.session_state.phase = 'milestone'
                st.rerun()
        with col_b:
            if st.button("← Back to tracking phase"):
                st.session_state.phase = 'tracking'
                st.rerun()


if __name__ == '__main__':
    main()
